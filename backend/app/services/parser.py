"""
Resume Parser Service — Sprint 2
Extracts structured data from PDF and DOCX files using PyMuPDF, pdfplumber, and python-docx.
"""
import re
import io
from typing import Any
import structlog

log = structlog.get_logger()


class ResumeParserService:
    """Parses resume files and extracts structured information."""

    # ─── Common skill keywords for extraction ────────────────────────────────
    SKILL_KEYWORDS = {
        "languages": ["python", "javascript", "typescript", "java", "c++", "c#", "go", "rust",
                      "ruby", "php", "swift", "kotlin", "scala", "r", "matlab", "sql", "html", "css"],
        "frameworks": ["react", "vue", "angular", "next.js", "nuxt", "fastapi", "django", "flask",
                       "spring", "express", "nest.js", "laravel", "rails", "pytorch", "tensorflow"],
        "databases": ["postgresql", "mysql", "mongodb", "redis", "sqlite", "elasticsearch",
                      "cassandra", "dynamodb", "firestore", "neo4j"],
        "cloud": ["aws", "gcp", "azure", "docker", "kubernetes", "terraform", "ansible",
                  "jenkins", "github actions", "circleci"],
        "tools": ["git", "linux", "nginx", "graphql", "rest api", "grpc", "kafka", "rabbitmq"],
        "ai_ml": ["machine learning", "deep learning", "nlp", "computer vision", "llm",
                  "scikit-learn", "pandas", "numpy", "openai", "langchain"],
    }

    async def parse(self, content: bytes, file_type: str) -> "ParsedResumeData":
        """Parse resume content and return structured data."""
        from app.schemas.resume import ParsedResumeData

        try:
            if file_type == "pdf":
                text = self._extract_pdf_text(content)
            elif file_type in ("docx", "doc"):
                text = self._extract_docx_text(content)
            else:
                text = content.decode("utf-8", errors="ignore")

            parsed = self._extract_entities(text)
            parsed["raw_text"] = text
            return ParsedResumeData(**parsed)

        except Exception as e:
            log.error("Resume parsing failed", error=str(e), file_type=file_type)
            from app.schemas.resume import ParsedResumeData
            return ParsedResumeData()

    def _extract_pdf_text(self, content: bytes) -> str:
        """Extract text from PDF using PyMuPDF with pdfplumber fallback."""
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(stream=content, filetype="pdf")
            text_parts = []
            for page in doc:
                text_parts.append(page.get_text("text"))
            doc.close()
            text = "\n".join(text_parts)
            if len(text.strip()) > 100:
                return text
        except Exception as e:
            log.warning("PyMuPDF extraction failed, falling back to pdfplumber", error=str(e))

        try:
            import pdfplumber
            with pdfplumber.open(io.BytesIO(content)) as pdf:
                text_parts = []
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
                return "\n".join(text_parts)
        except Exception as e:
            log.error("pdfplumber extraction failed", error=str(e))
            return ""

    def _extract_docx_text(self, content: bytes) -> str:
        """Extract text from DOCX using python-docx."""
        from docx import Document
        doc = Document(io.BytesIO(content))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text)
        return "\n".join(paragraphs)

    def _extract_entities(self, text: str) -> dict[str, Any]:
        """Extract structured entities from resume text."""
        return {
            "name": self._extract_name(text),
            "email": self._extract_email(text),
            "phone": self._extract_phone(text),
            "skills": self._extract_skills(text),
            "experience": self._extract_experience(text),
            "education": self._extract_education(text),
            "certifications": self._extract_certifications(text),
            "projects": self._extract_projects(text),
            "summary": self._extract_summary(text),
        }

    def _extract_email(self, text: str) -> str | None:
        pattern = r'\b[A-Za-z0-9._%+\-]+@[A-Za-z0-9.\-]+\.[A-Za-z]{2,}\b'
        match = re.search(pattern, text)
        return match.group(0) if match else None

    def _extract_phone(self, text: str) -> str | None:
        pattern = r'(\+?\d{1,3}[\s\-.]?)?\(?\d{3}\)?[\s\-.]?\d{3}[\s\-.]?\d{4}'
        match = re.search(pattern, text)
        return match.group(0).strip() if match else None

    def _extract_name(self, text: str) -> str | None:
        """Extract name from the top of the resume (heuristic)."""
        lines = [l.strip() for l in text.split("\n") if l.strip()]
        for line in lines[:5]:
            # Name is usually 2-4 words, no special chars, not all caps keywords
            words = line.split()
            if 2 <= len(words) <= 4:
                if all(w[0].isupper() or w.isupper() for w in words if w):
                    if not any(kw in line.lower() for kw in ["resume", "curriculum", "vitae", "cv"]):
                        return line
        return lines[0] if lines else None

    def _extract_skills(self, text: str) -> list[str]:
        """Extract skills using keyword matching across the full text."""
        text_lower = text.lower()
        found = set()

        for category, keywords in self.SKILL_KEYWORDS.items():
            for kw in keywords:
                if kw in text_lower:
                    found.add(kw.title() if len(kw) <= 3 else kw.capitalize())

        # Also extract from "Skills" section
        skills_section = self._extract_section(text, ["skills", "technical skills", "core competencies"])
        if skills_section:
            # Split by common delimiters
            raw_skills = re.split(r'[,|•·\n\t]+', skills_section)
            for s in raw_skills:
                cleaned = s.strip().strip("•-–")
                if 2 <= len(cleaned) <= 40 and not cleaned.isdigit():
                    found.add(cleaned)

        return sorted(list(found))[:50]  # Cap at 50 skills

    def _extract_experience(self, text: str) -> list[dict]:
        """Extract work experience entries."""
        section = self._extract_section(text, ["experience", "work experience", "employment"])
        if not section:
            return []

        entries = []
        # Split by year patterns as job boundaries
        job_blocks = re.split(r'\n(?=\d{4}|\w+ \d{4})', section)
        for block in job_blocks[:10]:  # Max 10 entries
            if len(block.strip()) > 20:
                lines = [l.strip() for l in block.split("\n") if l.strip()]
                entry = {
                    "title": lines[0] if lines else "",
                    "company": lines[1] if len(lines) > 1 else "",
                    "duration": self._extract_date_range(block),
                    "description": "\n".join(lines[2:5]) if len(lines) > 2 else "",
                }
                entries.append(entry)
        return entries

    def _extract_education(self, text: str) -> list[dict]:
        """Extract education entries."""
        section = self._extract_section(text, ["education", "academic background", "qualifications"])
        if not section:
            return []

        entries = []
        degree_patterns = [
            r'(?:b\.?tech|b\.?e|b\.?sc|m\.?tech|m\.?e|m\.?sc|mba|phd|bachelor|master|doctor)',
        ]
        lines = [l.strip() for l in section.split("\n") if l.strip()]
        for i, line in enumerate(lines[:8]):
            if any(re.search(p, line.lower()) for p in degree_patterns):
                entries.append({
                    "degree": line,
                    "institution": lines[i + 1] if i + 1 < len(lines) else "",
                    "year": self._extract_year(line + " " + (lines[i + 1] if i + 1 < len(lines) else "")),
                })
        return entries

    def _extract_certifications(self, text: str) -> list[str]:
        """Extract certifications."""
        section = self._extract_section(text, ["certifications", "certificates", "credentials"])
        if not section:
            return []
        items = [l.strip().lstrip("•-–") for l in section.split("\n") if l.strip()]
        return [i for i in items if 5 <= len(i) <= 150][:10]

    def _extract_projects(self, text: str) -> list[dict]:
        """Extract project entries."""
        section = self._extract_section(text, ["projects", "personal projects", "portfolio"])
        if not section:
            return []
        lines = [l.strip() for l in section.split("\n") if l.strip()]
        entries = []
        for i in range(0, min(len(lines), 15), 3):
            entries.append({
                "name": lines[i] if i < len(lines) else "",
                "description": lines[i + 1] if i + 1 < len(lines) else "",
            })
        return entries[:5]

    def _extract_summary(self, text: str) -> str | None:
        """Extract professional summary."""
        section = self._extract_section(text, ["summary", "profile", "objective", "about me"])
        if section:
            return section[:500].strip()
        return None

    def _extract_section(self, text: str, headings: list[str]) -> str | None:
        """Extract a named section from the resume text."""
        lines = text.split("\n")
        section_start = -1
        pattern = "|".join(re.escape(h) for h in headings)

        for i, line in enumerate(lines):
            if re.search(pattern, line.lower()):
                section_start = i + 1
                break

        if section_start == -1:
            return None

        section_lines = []
        for line in lines[section_start:]:
            # Stop at the next section heading (all caps or ends with colon)
            if re.match(r'^[A-Z][A-Z\s]+:?$', line.strip()) and len(line.strip()) > 3:
                break
            section_lines.append(line)
            if len(section_lines) > 50:
                break

        return "\n".join(section_lines).strip()

    def _extract_date_range(self, text: str) -> str | None:
        pattern = r'(\w+ \d{4}|\d{4})\s*[-–to]+\s*(\w+ \d{4}|\d{4}|present|current)'
        match = re.search(pattern, text, re.IGNORECASE)
        return match.group(0) if match else None

    def _extract_year(self, text: str) -> str | None:
        match = re.search(r'\b(19|20)\d{2}\b', text)
        return match.group(0) if match else None
